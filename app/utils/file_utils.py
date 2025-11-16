"""File utility functions."""

import hashlib
import io
import re
from typing import Optional, Tuple

import fitz  # PyMuPDF
import magic
from langchain_community.document_loaders import Blob
from langchain_community.document_loaders.parsers.generic import MimeTypeBasedParser
from langchain_community.document_loaders.parsers.html.bs4 import BS4HTMLParser
from langchain_community.document_loaders.parsers.msword import MsWordParser
from langchain_community.document_loaders.parsers.txt import TextParser
from langchain_core.documents import Document

try:
    from striprtf.striprtf import rtf_to_text

    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.core.logging import logger
from app.exceptions.custom_exceptions import (
    FileProcessingError,
    FileSizeLimitError,
    UnsupportedFileTypeError,
)


class PyMuPDFParser:
    """Advanced PDF parser using PyMuPDF for block-based text extraction."""

    def parse(self, blob: Blob) -> list[Document]:
        """Parse PDF blob into documents with block-based extraction.

        Args:
            blob: The PDF blob to parse

        Returns:
            List of documents
        """
        return list(self.lazy_parse(blob))

    def lazy_parse(self, blob: Blob):
        """Lazy parse PDF blob into documents with block-based extraction.

        Args:
            blob: The PDF blob to parse

        Yields:
            Documents with structured text blocks
        """
        try:
            # Open PDF from bytes
            pdf_bytes = blob.as_bytes()
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")

            # Extract text blocks from all pages
            all_blocks = []

            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]

                # Get text blocks with position information
                # blocks format: (x0, y0, x1, y1, "text", block_no, block_type)
                blocks = page.get_text("blocks")

                for block in blocks:
                    # block[4] is the text content
                    # block[5] is the block number
                    # block[6] is the block type (0=text, 1=image)

                    if block[6] == 0:  # Only process text blocks
                        text = block[4].strip()
                        if text:  # Only add non-empty blocks
                            # Calculate block position for ordering
                            x0, y0, x1, y1 = block[0], block[1], block[2], block[3]

                            all_blocks.append(
                                {
                                    "text": text,
                                    "page": page_num + 1,
                                    "x0": x0,
                                    "y0": y0,
                                    "x1": x1,
                                    "y1": y1,
                                    "block_no": block[5],
                                }
                            )

            pdf_document.close()

            # Sort blocks by page and vertical position (top to bottom)
            all_blocks.sort(key=lambda b: (b["page"], b["y0"], b["x0"]))

            # Group blocks into meaningful sections
            structured_text = self._structure_blocks(all_blocks)

            yield Document(
                page_content=structured_text,
                metadata={
                    "source": blob.path,
                    "total_blocks": len(all_blocks),
                    "parser": "PyMuPDF",
                },
            )

        except Exception as e:
            logger.error(f"PyMuPDF parsing failed: {str(e)}")
            # Yield error document instead of raising
            yield Document(
                page_content=f"Failed to parse PDF: {str(e)}",
                metadata={"source": blob.path, "error": True},
            )

    def _structure_blocks(self, blocks: list[dict]) -> str:
        """Structure text blocks into meaningful content.

        This method organizes blocks preserving their natural layout
        while maintaining readability.

        Args:
            blocks: List of text blocks with position info

        Returns:
            Structured text content
        """
        if not blocks:
            return ""

        structured_lines = []
        current_page = None

        for block in blocks:
            # Add page separator if new page
            if current_page != block["page"]:
                if current_page is not None:
                    structured_lines.append("")  # Empty line between pages
                current_page = block["page"]

            # Clean and add block text
            text = block["text"]

            # Remove excessive whitespace within block but preserve structure
            text = re.sub(r"[ \t]+", " ", text)  # Multiple spaces to single space
            text = re.sub(r"\n{3,}", "\n\n", text)  # Max 2 newlines

            # Add the block
            structured_lines.append(text)

        # Join all blocks with single newlines
        # Blocks already have internal structure, so we just separate them
        result = "\n".join(structured_lines)

        # Final cleanup
        result = re.sub(r"\n{3,}", "\n\n", result)  # Max 2 consecutive newlines

        return result.strip()


class CustomWordParser:
    """Custom Word document parser with fallback mechanisms."""

    def parse(self, blob: Blob) -> list[Document]:
        """Parse Word document blob into documents.

        Args:
            blob: The blob to parse

        Returns:
            List of documents
        """
        return list(self.lazy_parse(blob))

    def lazy_parse(self, blob: Blob):
        """Lazy parse Word document blob into documents.

        Args:
            blob: The blob to parse

        Yields:
            Documents
        """
        try:
            # First try direct python-docx parsing
            if DOCX_AVAILABLE:
                content = self._extract_with_python_docx(blob.as_bytes())
                if content:
                    yield Document(page_content=content, metadata={"source": blob.path})
                    return

            # Fallback to original MsWordParser
            try:
                ms_parser = MsWordParser()
                yield from ms_parser.lazy_parse(blob)
                return
            except Exception as fallback_error:
                logger.warning(
                    f"MsWordParser failed: {fallback_error}, trying text extraction"
                )

                # Last resort: try to extract as much text as possible
                content = self._extract_text_fallback(blob.as_bytes())
                yield Document(page_content=content, metadata={"source": blob.path})

        except Exception as e:
            logger.error(f"Failed to parse Word document: {str(e)}")
            # Don't raise error, just yield a document with error message
            yield Document(
                page_content=f"Failed to parse Word document: {str(e)}",
                metadata={"source": blob.path, "error": True},
            )

    def _extract_with_python_docx(self, content: bytes) -> str:
        """Extract text using python-docx library directly.

        Args:
            content: Document content as bytes

        Returns:
            Extracted text
        """
        try:
            doc = DocxDocument(io.BytesIO(content))

            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())

            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_cells.append(cell_text)
                    if row_cells:
                        table_rows.append(" | ".join(row_cells))
                if table_rows:
                    tables_text.extend(table_rows)

            # Combine all text
            all_text = []
            all_text.extend(paragraphs)
            if tables_text:
                all_text.append("\n--- Tables ---")
                all_text.extend(tables_text)

            return "\n".join(all_text)

        except Exception as e:
            logger.warning(f"python-docx extraction failed: {e}")
            return ""

    def _extract_text_fallback(self, content: bytes) -> str:
        """Fallback text extraction for problematic Word documents.

        Args:
            content: Document content as bytes

        Returns:
            Extracted text (may be partial)
        """
        try:
            # Try to decode as much text as possible from the binary content
            # This is a very basic fallback that looks for text patterns

            # Convert to string with error handling
            text_parts = []

            # Try different encodings
            for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
                try:
                    decoded = content.decode(encoding, errors="ignore")
                    # Look for readable text (letters, numbers, common punctuation)
                    readable_text = re.findall(
                        r"[a-zA-Z0-9\s\.\,\;\:\!\?\-\(\)]+", decoded
                    )
                    if readable_text:
                        text_parts.extend(
                            [t.strip() for t in readable_text if len(t.strip()) > 2]
                        )
                        break
                except Exception:
                    continue

            if text_parts:
                # Clean and deduplicate
                unique_parts = list(
                    dict.fromkeys(text_parts)
                )  # Remove duplicates while preserving order
                return "\n".join(unique_parts[:100])  # Limit to first 100 parts

            return "Failed to extract readable text from document"

        except Exception as e:
            logger.error(f"Fallback text extraction failed: {e}")
            return "Document text extraction failed"


class RTFParser:
    """Custom RTF parser for handling RTF files."""

    def parse(self, blob: Blob) -> list[Document]:
        """Parse RTF blob into documents.

        Args:
            blob: The blob to parse

        Returns:
            List of documents
        """
        return list(self.lazy_parse(blob))

    def lazy_parse(self, blob: Blob):
        """Lazy parse RTF blob into documents.

        Args:
            blob: The blob to parse

        Yields:
            Documents
        """
        try:
            if RTF_AVAILABLE:
                # Use striprtf for proper RTF parsing
                rtf_content = blob.as_string()
                text_content = rtf_to_text(rtf_content)
            else:
                # Fallback: simple RTF to text conversion
                rtf_content = blob.as_string()
                text_content = self._simple_rtf_to_text(rtf_content)

            yield Document(page_content=text_content, metadata={"source": blob.path})

        except Exception as e:
            logger.error(f"Failed to parse RTF: {str(e)}")
            # Don't raise error, just yield a document with error message
            yield Document(
                page_content=f"Failed to parse RTF file: {str(e)}",
                metadata={"source": blob.path, "error": True},
            )

    def _simple_rtf_to_text(self, rtf_content: str) -> str:
        """Simple RTF to text conversion fallback.

        Args:
            rtf_content: RTF content as string

        Returns:
            Extracted text
        """
        # Remove RTF control words and groups
        text = re.sub(r"\\[a-z]+\d*", "", rtf_content)
        text = re.sub(r"[{}]", "", text)
        text = re.sub(r"\\[^a-z]", "", text)

        # Clean up whitespace
        text = re.sub(r"\s+", " ", text)
        text = text.strip()

        return text


class FileProcessor:
    """Utility class for file processing operations."""

    def __init__(self, max_file_size_mb: int = 10):
        """Initialize file processor.

        Args:
            max_file_size_mb: Maximum file size in MB
        """
        self.max_file_size_mb = max_file_size_mb

        # Initialize custom parsers
        self.rtf_parser = RTFParser()
        self.word_parser = CustomWordParser()
        self.pymupdf_parser = PyMuPDFParser()  # PyMuPDF for better PDF extraction

        # Define supported handlers
        self.SUPPORTED_HANDLERS = {
            # PDF formats - using PyMuPDF for block-based extraction
            "application/pdf": self.pymupdf_parser,
            # Text formats
            "text/plain": TextParser(),
            "text/plain; charset=utf-8": TextParser(),
            "text/plain; charset=ascii": TextParser(),
            "application/octet-stream": TextParser(),  # Sometimes .txt files are detected as this
            # HTML formats
            "text/html": BS4HTMLParser(),
            "text/html; charset=utf-8": BS4HTMLParser(),
            "application/xhtml+xml": BS4HTMLParser(),
            # Word document formats - using custom parser for better compatibility
            "application/msword": self.word_parser,
            "application/doc": self.word_parser,
            "application/vnd.ms-word": self.word_parser,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self.word_parser,
            "application/vnd.ms-word.document.macroEnabled.12": self.word_parser,
            # RTF formats
            "text/rtf": self.rtf_parser,
            "application/rtf": self.rtf_parser,
            # Additional text-based formats
            "text/csv": TextParser(),
            "application/csv": TextParser(),
            "text/tab-separated-values": TextParser(),
            "application/xml": BS4HTMLParser(),  # XML can be parsed like HTML
            "text/xml": BS4HTMLParser(),
            # Image formats (processed via OpenAI Vision API)
            "image/jpeg": None,  # Vision API
            "image/jpg": None,  # Vision API
            "image/png": None,  # Vision API
            "image/webp": None,  # Vision API
            "image/gif": None,  # Vision API (non-animated)
        }

        self.supported_mimetypes = sorted(self.SUPPORTED_HANDLERS.keys())
        self.parser = MimeTypeBasedParser(
            handlers=self.SUPPORTED_HANDLERS, fallback_parser=None
        )

        try:
            self.mime = magic.Magic(mime=True)
        except Exception as e:
            logger.error(f"Failed to initialize magic library: {str(e)}")
            raise FileProcessingError("Failed to initialize file type detection") from e

    def get_content_hash(self, content: bytes) -> str:
        """Generate SHA256 hash of content for caching.

        Args:
            content: File content as bytes

        Returns:
            Hexadecimal hash string
        """
        return hashlib.sha256(content).hexdigest()

    def guess_mimetype(self, file_bytes: bytes, filename: str = "") -> str:
        """Detect MIME type of file content with filename fallback.

        Args:
            file_bytes: File content as bytes
            filename: Original filename for extension-based fallback

        Returns:
            MIME type string

        Raises:
            UnsupportedFileTypeError: If file type is not supported
            FileProcessingError: If detection fails
        """
        try:
            # Try magic detection first
            mime_type = self.mime.from_buffer(file_bytes)

            # If detected type is supported, use it
            if mime_type in self.supported_mimetypes:
                return mime_type

            # Try filename extension fallback
            if filename:
                extension_mime = self._guess_mimetype_from_extension(filename)
                if extension_mime and extension_mime in self.supported_mimetypes:
                    logger.info(
                        f"Using extension-based MIME type: {extension_mime} for {filename}"
                    )
                    return extension_mime

            # If still not supported, raise error
            raise UnsupportedFileTypeError(
                f"Unsupported file type: {mime_type}. "
                f"Supported types: {', '.join(self.supported_mimetypes)}"
            )

        except UnsupportedFileTypeError:
            raise
        except Exception as e:
            logger.error(f"Failed to detect MIME type: {str(e)}")
            raise FileProcessingError(f"Failed to detect file type: {str(e)}") from e

    def is_image_format(self, mime_type: str) -> bool:
        """Check if MIME type is an image format.

        Args:
            mime_type: MIME type string

        Returns:
            True if image format, False otherwise
        """
        return mime_type.startswith("image/")

    def _guess_mimetype_from_extension(self, filename: str) -> Optional[str]:
        """Guess MIME type from file extension.

        Args:
            filename: Original filename

        Returns:
            MIME type string or None if unknown
        """
        if not filename or "." not in filename:
            return None

        extension = filename.lower().split(".")[-1]

        extension_mapping = {
            "pdf": "application/pdf",
            "txt": "text/plain",
            "html": "text/html",
            "htm": "text/html",
            "doc": "application/msword",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "rtf": "text/rtf",
            "csv": "text/csv",
            "xml": "application/xml",
            "xhtml": "application/xhtml+xml",
            # Image formats
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
            "png": "image/png",
            "webp": "image/webp",
            "gif": "image/gif",
        }

        return extension_mapping.get(extension)

    def validate_file_size(self, content: bytes) -> None:
        """Validate file size against maximum limit.

        Args:
            content: File content as bytes

        Raises:
            FileSizeLimitError: If file exceeds size limit
        """
        size_mb = len(content) / (1024 * 1024)
        if size_mb > self.max_file_size_mb:
            raise FileSizeLimitError(
                f"File size ({size_mb:.2f}MB) exceeds maximum limit of "
                f"{self.max_file_size_mb}MB"
            )

    def extract_text_from_content(
        self, content: bytes, filename: str
    ) -> Tuple[str, str]:
        """Extract text from file content.

        Args:
            content: File content as bytes
            filename: Original filename

        Returns:
            Tuple of (extracted_text, mime_type)
            For image files, extracted_text will be empty string

        Raises:
            FileProcessingError: If extraction fails
        """
        try:
            # Validate file size
            self.validate_file_size(content)

            # Detect MIME type with filename fallback
            mime_type = self.guess_mimetype(content, filename)

            # Check if it's an image format - images are processed via Vision API
            if self.is_image_format(mime_type):
                logger.info(
                    f"Image file detected: {filename} ({mime_type}) - will use Vision API"
                )
                return "", mime_type  # Empty text, Vision API will handle it

            # Create blob for processing
            blob = Blob.from_data(data=content, path=filename, mime_type=mime_type)

            # Extract text
            documents = self.parser.parse(blob)
            extracted_text = "\n".join([doc.page_content for doc in documents])

            # Clean text
            cleaned_text = self.clean_text(extracted_text)

            logger.info(f"Successfully extracted text from {filename} ({mime_type})")
            return cleaned_text, mime_type

        except (UnsupportedFileTypeError, FileSizeLimitError):
            raise
        except Exception as e:
            logger.error(f"Failed to extract text from {filename}: {str(e)}")
            raise FileProcessingError(f"Failed to process file: {str(e)}") from e

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean extracted text by removing HTML tags and extra whitespace.

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text
        """
        # Remove HTML tags
        clean = re.compile("<.*?>")
        text = re.sub(clean, "", text)

        # Remove multiple newlines
        text = re.sub(r"\n+", "\n", text)

        # Remove multiple quotes
        text = re.sub(r'"+', '"', text)

        # Strip whitespace
        text = text.strip()

        return text

    def get_supported_formats(self) -> list:
        """Get list of supported MIME types.

        Returns:
            List of supported MIME types
        """
        return self.supported_mimetypes
